from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from core.cv_formatters import contact_parts, language_strings, skill_strings
from core.i18n import tr


def _set_layout(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.8)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def _p(doc, text, bold=False, size=9.8, italic=False, after=1):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(after)


def _section(doc, title):
    _p(doc, title, bold=True, size=10.8, after=2)


def _structured_line(title, subtitle, period):
    return " | ".join([x for x in [title.strip(), subtitle.strip(), period.strip()] if x])


def build_compact_cv(profile):
    doc = Document()
    _set_layout(doc)

    lang = profile.language

    header = profile.name
    if profile.job_title.strip():
        header += f" - {profile.job_title.strip()}"
    _p(doc, header, bold=True, size=16, after=2)

    parts = contact_parts(profile)
    if parts:
        _p(doc, " | ".join(parts), size=9.3, after=3)

    if profile.summary.strip():
        _section(doc, tr(lang, "summary"))
        _p(doc, profile.summary.strip(), after=2)

    if profile.experience:
        _section(doc, tr(lang, "experience"))
        for entry in profile.experience:
            line = " | ".join(
                [x for x in [entry.title.strip(), entry.company.strip(), entry.location.strip(), entry.period.strip()] if x]
            )
            if line:
                _p(doc, line, bold=True, size=9.9, after=2)
            for detail in entry.details:
                _p(doc, f"- {detail}", size=9.6, after=1)

    if profile.education:
        _section(doc, tr(lang, "education"))
        for entry in profile.education:
            line = " | ".join(
                [x for x in [entry.degree.strip(), entry.school.strip(), entry.location.strip(), entry.period.strip()] if x]
            )
            if line:
                _p(doc, line, bold=True, size=9.9, after=2)
            for detail in entry.details:
                _p(doc, f"- {detail}", size=9.6, after=1)

    if profile.skills:
        _section(doc, tr(lang, "skills"))
        _p(doc, " | ".join(skill_strings(profile.skills)), size=9.7, after=1)

    if profile.languages:
        _section(doc, tr(lang, "languages"))
        _p(doc, " | ".join(language_strings(profile.languages)), size=9.7, after=1)

    if profile.projects:
        _section(doc, tr(lang, "projects"))
        for project in profile.projects:
            line = _structured_line(
                getattr(project, "title", str(project)),
                getattr(project, "organization", ""),
                getattr(project, "period", ""),
            )
            if line:
                _p(doc, line, bold=True, size=9.9, after=1)
            summary = str(getattr(project, "summary", "") or "").strip()
            if summary:
                _p(doc, summary, size=9.6, after=1)

    if profile.certificates:
        _section(doc, tr(lang, "certificates"))
        for certificate in profile.certificates:
            line = _structured_line(
                getattr(certificate, "title", str(certificate)),
                getattr(certificate, "issuer", ""),
                getattr(certificate, "period", ""),
            )
            if line:
                _p(doc, line, bold=True, size=9.9, after=1)
            summary = str(getattr(certificate, "summary", "") or "").strip()
            if summary:
                _p(doc, summary, size=9.6, after=1)

    return doc

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from core.cv_formatters import contact_parts, language_strings, skill_strings
from core.i18n import tr


def _set_layout(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _text(doc, text, bold=False, size=11, after=2):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(after)


def _section(doc, title):
    _text(doc, title.upper(), bold=True, size=11.5, after=3)


def _project_like_block(doc, title, subtitle, period, summary):
    title_line = " | ".join([x for x in [title.strip(), subtitle.strip()] if x])
    if title_line:
        _text(doc, title_line, bold=True, size=11, after=2)
    if period.strip():
        _text(doc, period.strip(), size=10.5, after=2)
    if summary.strip():
        _text(doc, f"- {summary.strip()}", size=10.8, after=1)


def build_ats_cv(profile):
    doc = Document()
    _set_layout(doc)

    lang = profile.language

    _text(doc, profile.name, bold=True, size=18, after=1)

    if profile.job_title.strip():
        _text(doc, profile.job_title.strip(), size=11.5, after=2)

    parts = contact_parts(profile)
    if parts:
        _text(doc, " | ".join(parts), size=10.5, after=4)

    if profile.summary.strip():
        _section(doc, tr(lang, "summary"))
        _text(doc, profile.summary.strip(), size=11, after=3)

    if profile.experience:
        _section(doc, tr(lang, "experience"))
        for entry in profile.experience:
            title_line = " | ".join([x for x in [entry.title.strip(), entry.company.strip(), entry.location.strip()] if x])
            if title_line:
                _text(doc, title_line, bold=True, size=11, after=2)
            if entry.period.strip():
                _text(doc, entry.period.strip(), size=10.5, after=2)
            for detail in entry.details:
                _text(doc, f"- {detail}", size=10.8, after=1)

    if profile.education:
        _section(doc, tr(lang, "education"))
        for entry in profile.education:
            degree_line = " | ".join([x for x in [entry.degree.strip(), entry.school.strip(), entry.location.strip()] if x])
            if degree_line:
                _text(doc, degree_line, bold=True, size=11, after=2)
            if entry.period.strip():
                _text(doc, entry.period.strip(), size=10.5, after=2)
            for detail in entry.details:
                _text(doc, f"- {detail}", size=10.8, after=1)

    if profile.skills:
        _section(doc, tr(lang, "skills"))
        for item in skill_strings(profile.skills):
            _text(doc, f"- {item}", size=10.8, after=1)

    if profile.languages:
        _section(doc, tr(lang, "languages"))
        for item in language_strings(profile.languages):
            _text(doc, f"- {item}", size=10.8, after=1)

    if profile.projects:
        _section(doc, tr(lang, "projects"))
        for entry in profile.projects:
            _project_like_block(
                doc,
                getattr(entry, "title", str(entry)),
                getattr(entry, "organization", ""),
                getattr(entry, "period", ""),
                getattr(entry, "summary", ""),
            )

    if profile.certificates:
        _section(doc, tr(lang, "certificates"))
        for entry in profile.certificates:
            _project_like_block(
                doc,
                getattr(entry, "title", str(entry)),
                getattr(entry, "issuer", ""),
                getattr(entry, "period", ""),
                getattr(entry, "summary", ""),
            )

    return doc

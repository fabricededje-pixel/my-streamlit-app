from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from core.cv_formatters import contact_parts, language_strings, skill_strings
from core.i18n import tr


def _set_layout(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def _add_separator(doc, color="4F81BD"):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)

    paragraph_props = paragraph._element.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    border.append(bottom)
    paragraph_props.append(border)


def _add_heading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(13.5)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)


def _add_text(doc, text, bold=False, size=10.5, italic=False, space_after=2):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(space_after)


def _add_bullets(doc, items):
    for item in items:
        if item:
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(item)
            run.font.size = Pt(10.5)
            paragraph.paragraph_format.space_after = Pt(2)


def _add_experience(doc, entry):
    if entry.title.strip():
        _add_text(doc, entry.title.strip(), bold=True, size=11.5, space_after=2)

    line2 = " | ".join([x for x in [entry.company.strip(), entry.location.strip()] if x])
    if line2:
        _add_text(doc, line2, size=10.5, italic=True, space_after=2)

    if entry.period.strip():
        _add_text(doc, entry.period.strip(), size=10.2, space_after=2)

    _add_bullets(doc, entry.details)


def _add_education(doc, entry):
    if entry.degree.strip():
        _add_text(doc, entry.degree.strip(), bold=True, size=11.5, space_after=2)

    line2 = " | ".join([x for x in [entry.school.strip(), entry.location.strip()] if x])
    if line2:
        _add_text(doc, line2, size=10.5, italic=True, space_after=2)

    if entry.period.strip():
        _add_text(doc, entry.period.strip(), size=10.2, space_after=2)

    _add_bullets(doc, entry.details)


def _add_project_like(doc, title, subtitle, period, summary):
    if title.strip():
        _add_text(doc, title.strip(), bold=True, size=11.2, space_after=2)
    meta = " | ".join([x for x in [subtitle.strip(), period.strip()] if x])
    if meta:
        _add_text(doc, meta, size=10.2, italic=True, space_after=2)
    if summary.strip():
        _add_text(doc, summary.strip(), size=10.4, space_after=3)


def build_classic_cv(profile):
    doc = Document()
    _set_layout(doc)

    lang = profile.language

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(profile.name)
    run.bold = True
    run.font.size = Pt(22)
    paragraph.paragraph_format.space_after = Pt(2)

    if profile.job_title.strip():
        _add_text(doc, profile.job_title.strip(), size=12, space_after=3)

    parts = contact_parts(profile)
    if parts:
        _add_text(doc, " | ".join(parts), size=10.2, space_after=5)

    _add_separator(doc)

    if profile.summary.strip():
        _add_heading(doc, tr(lang, "summary"))
        _add_text(doc, profile.summary.strip(), size=10.5, space_after=3)

    if profile.experience:
        _add_heading(doc, tr(lang, "experience"))
        for entry in profile.experience:
            _add_experience(doc, entry)

    if profile.education:
        _add_heading(doc, tr(lang, "education"))
        for entry in profile.education:
            _add_education(doc, entry)

    if profile.skills:
        _add_heading(doc, tr(lang, "skills"))
        _add_bullets(doc, skill_strings(profile.skills))

    if profile.languages:
        _add_heading(doc, tr(lang, "languages"))
        _add_bullets(doc, language_strings(profile.languages))

    if profile.projects:
        _add_heading(doc, tr(lang, "projects"))
        for entry in profile.projects:
            _add_project_like(
                doc,
                getattr(entry, "title", str(entry)),
                getattr(entry, "organization", ""),
                getattr(entry, "period", ""),
                getattr(entry, "summary", ""),
            )

    if profile.certificates:
        _add_heading(doc, tr(lang, "certificates"))
        for entry in profile.certificates:
            _add_project_like(
                doc,
                getattr(entry, "title", str(entry)),
                getattr(entry, "issuer", ""),
                getattr(entry, "period", ""),
                getattr(entry, "summary", ""),
            )

    return doc

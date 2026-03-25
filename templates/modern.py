from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from core.cv_formatters import contact_parts, language_strings, skill_strings
from core.i18n import tr


DEFAULT_MODERN_COLORS = {
    "header_fill": "2F607D",
    "sidebar_fill": "E9EDF2",
    "header_line": "D7E1EA",
    "text_dark": "1F2933",
    "text_muted": "5F6C7B",
}


def _set_layout(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _set_cell_shading(cell, fill):
    cell_props = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    cell_props.append(shade)


def _remove_table_borders(table):
    table_props = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    table_props.append(borders)


def _set_cell_padding(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _paragraph(cell, text="", *, bold=False, size=10, color, after=2, upper=False):
    paragraph = cell.add_paragraph()
    run = paragraph.add_run(text.upper() if upper else text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    paragraph.paragraph_format.space_after = Pt(after)
    return paragraph


def _first_paragraph(cell, text="", *, bold=False, size=10, color, after=2):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    paragraph.paragraph_format.space_after = Pt(after)
    return paragraph


def _section_title(cell, text, colors):
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(colors["text_dark"])
    return paragraph


def _sidebar_title(cell, text, colors):
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(colors["text_dark"])
    return paragraph


def _sidebar_lines(cell, items, colors):
    for item in items:
        if str(item).strip():
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(str(item))
            run.font.name = "Arial"
            run.font.size = Pt(9.3)
            run.font.color.rgb = RGBColor.from_string(colors["text_dark"])


def _sidebar_bullets(cell, items, colors):
    for item in items:
        if str(item).strip():
            paragraph = cell.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(1.5)
            run = paragraph.add_run(str(item))
            run.font.name = "Arial"
            run.font.size = Pt(9.1)
            run.font.color.rgb = RGBColor.from_string(colors["text_dark"])


def _entry_block(cell, title, subtitle, period, details, colors):
    title_line = title.strip()
    if period.strip():
        title_line = f"{title_line}, {period.strip()}" if title_line else period.strip()
    if title_line:
        _paragraph(cell, title_line, bold=True, size=10.6, color=colors["text_dark"], after=1)
    if subtitle.strip():
        _paragraph(cell, subtitle, bold=True, size=10.1, color=colors["text_muted"], after=2)
    for detail in details:
        if str(detail).strip():
            paragraph = cell.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(1.5)
            run = paragraph.add_run(str(detail))
            run.font.name = "Arial"
            run.font.size = Pt(9.6)
            run.font.color.rgb = RGBColor.from_string(colors["text_dark"])


def _summary_entry_block(cell, title, subtitle, period, summary, colors):
    title_line = title.strip()
    if period.strip():
        title_line = f"{title_line}, {period.strip()}" if title_line else period.strip()
    if title_line:
        _paragraph(cell, title_line, bold=True, size=10.6, color=colors["text_dark"], after=1)
    if subtitle.strip():
        _paragraph(cell, subtitle, bold=True, size=10.1, color=colors["text_muted"], after=2)
    if summary.strip():
        _paragraph(cell, summary.strip(), size=9.8, color=colors["text_dark"], after=4)


def build_modern_cv(profile, colors=None):
    doc = Document()
    _set_layout(doc)
    lang = profile.language
    colors = colors or DEFAULT_MODERN_COLORS

    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    _remove_table_borders(header)

    left = header.cell(0, 0)
    right = header.cell(0, 1)
    left.width = Cm(4.4)
    right.width = Cm(13.6)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_shading(left, colors["header_fill"])
    _set_cell_shading(right, colors["header_fill"])
    _set_cell_padding(left, top=110, start=110, bottom=110, end=110)
    _set_cell_padding(right, top=110, start=130, bottom=110, end=110)

    if profile.photo_path:
        try:
            paragraph_img = _first_paragraph(left, "", color=colors["text_dark"], after=0)
            run_img = paragraph_img.add_run()
            run_img.add_picture(profile.photo_path, width=Cm(3.0))
        except Exception:
            _first_paragraph(left, "", color=colors["text_dark"], after=0)
    else:
        _first_paragraph(left, "", color=colors["text_dark"], after=0)

    _first_paragraph(right, profile.name.upper(), bold=True, size=22, color="FFFFFF", after=2)
    if profile.job_title.strip():
        _paragraph(right, profile.job_title.strip(), size=11, color="E6EEF5", after=0)

    separator = doc.add_paragraph()
    separator.paragraph_format.space_after = Pt(0)
    p_pr = separator._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), colors["header_line"])
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    body = doc.add_table(rows=1, cols=2)
    body.alignment = WD_TABLE_ALIGNMENT.CENTER
    body.autofit = False
    _remove_table_borders(body)

    sidebar = body.cell(0, 0)
    main = body.cell(0, 1)
    sidebar.width = Cm(4.4)
    main.width = Cm(13.6)
    sidebar.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    main.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_cell_shading(sidebar, colors["sidebar_fill"])
    _set_cell_padding(sidebar, top=120, start=100, bottom=100, end=100)
    _set_cell_padding(main, top=120, start=120, bottom=100, end=80)

    contact = contact_parts(profile)
    if contact:
        _sidebar_title(sidebar, "Kontakt", colors)
        _sidebar_lines(sidebar, contact, colors)

    if profile.skills:
        _sidebar_title(sidebar, tr(lang, "skills"), colors)
        _sidebar_lines(sidebar, skill_strings(profile.skills), colors)

    if profile.languages:
        _sidebar_title(sidebar, tr(lang, "languages"), colors)
        _sidebar_lines(sidebar, language_strings(profile.languages), colors)

    if profile.summary.strip():
        _section_title(main, tr(lang, "summary"), colors)
        _paragraph(main, profile.summary.strip(), size=10, color=colors["text_dark"], after=5)

    if profile.experience:
        _section_title(main, tr(lang, "experience"), colors)
        for entry in profile.experience:
            subtitle = " | ".join([x for x in [entry.company.strip(), entry.location.strip()] if x])
            _entry_block(main, entry.title, subtitle, entry.period, entry.details, colors)

    if profile.education:
        _section_title(main, tr(lang, "education"), colors)
        for entry in profile.education:
            subtitle = " | ".join([x for x in [entry.school.strip(), entry.location.strip()] if x])
            _entry_block(main, entry.degree, subtitle, entry.period, entry.details, colors)

    if profile.projects:
        _section_title(main, tr(lang, "projects"), colors)
        for entry in profile.projects:
            _summary_entry_block(
                main,
                str(getattr(entry, "title", entry)),
                getattr(entry, "organization", ""),
                getattr(entry, "period", ""),
                getattr(entry, "summary", ""),
                colors,
            )

    if profile.certificates:
        _section_title(main, tr(lang, "certificates"), colors)
        for entry in profile.certificates:
            _summary_entry_block(
                main,
                str(getattr(entry, "title", entry)),
                getattr(entry, "issuer", ""),
                getattr(entry, "period", ""),
                getattr(entry, "summary", ""),
                colors,
            )

    return doc

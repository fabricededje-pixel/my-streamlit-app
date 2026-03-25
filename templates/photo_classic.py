from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from core.cv_formatters import language_strings, skill_strings
from core.i18n import tr


def _set_layout(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def _remove_table_borders(table):
    table_props = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    table_props.append(borders)


def _separator(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph_props = paragraph._element.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4F81BD")
    border.append(bottom)
    paragraph_props.append(border)


def _heading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(13.5)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)


def _text(doc, text, bold=False, italic=False, size=10.5, after=2):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(after)


def _bullets(doc, items):
    for item in items:
        if item:
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(str(item))
            run.font.size = Pt(10.5)
            paragraph.paragraph_format.space_after = Pt(2)


def _add_header_with_photo(doc, profile):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _remove_table_borders(table)

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.width = Cm(12.5)
    right.width = Cm(4.0)

    paragraph = left.paragraphs[0]
    run = paragraph.add_run(profile.name)
    run.bold = True
    run.font.size = Pt(21)
    paragraph.paragraph_format.space_after = Pt(2)

    if profile.job_title.strip():
        paragraph_job = left.add_paragraph()
        run_job = paragraph_job.add_run(profile.job_title.strip())
        run_job.font.size = Pt(11.5)
        paragraph_job.paragraph_format.space_after = Pt(4)

    parts = [profile.email, profile.phone, profile.city, profile.linkedin]
    parts = [item for item in parts if item and str(item).strip()]
    if parts:
        paragraph_contact = left.add_paragraph()
        run_contact = paragraph_contact.add_run(" | ".join(parts))
        run_contact.font.size = Pt(10.2)

    if profile.photo_path:
        try:
            paragraph_image = right.paragraphs[0]
            run_image = paragraph_image.add_run()
            run_image.add_picture(profile.photo_path, width=Cm(3.2))
        except Exception:
            pass


def _exp(doc, entry):
    if entry.title.strip():
        _text(doc, entry.title.strip(), bold=True, size=11.5, after=2)

    second = " | ".join([x for x in [entry.company.strip(), entry.location.strip()] if x])
    if second:
        _text(doc, second, italic=True, size=10.5, after=2)

    if entry.period.strip():
        _text(doc, entry.period.strip(), size=10.2, after=2)

    _bullets(doc, entry.details)


def _edu(doc, entry):
    if entry.degree.strip():
        _text(doc, entry.degree.strip(), bold=True, size=11.5, after=2)

    second = " | ".join([x for x in [entry.school.strip(), entry.location.strip()] if x])
    if second:
        _text(doc, second, italic=True, size=10.5, after=2)

    if entry.period.strip():
        _text(doc, entry.period.strip(), size=10.2, after=2)

    _bullets(doc, entry.details)


def _project_like(doc, title, subtitle, period, summary):
    if title.strip():
        _text(doc, title.strip(), bold=True, size=11.2, after=2)
    meta = " | ".join([x for x in [subtitle.strip(), period.strip()] if x])
    if meta:
        _text(doc, meta, italic=True, size=10.3, after=2)
    if summary.strip():
        _text(doc, summary.strip(), size=10.4, after=3)


def build_photo_classic_cv(profile):
    doc = Document()
    _set_layout(doc)

    lang = profile.language

    _add_header_with_photo(doc, profile)
    _separator(doc)

    if profile.summary.strip():
        _heading(doc, tr(lang, "summary"))
        _text(doc, profile.summary.strip(), after=3)

    if profile.experience:
        _heading(doc, tr(lang, "experience"))
        for entry in profile.experience:
            _exp(doc, entry)

    if profile.education:
        _heading(doc, tr(lang, "education"))
        for entry in profile.education:
            _edu(doc, entry)

    if profile.skills:
        _heading(doc, tr(lang, "skills"))
        _bullets(doc, skill_strings(profile.skills))

    if profile.languages:
        _heading(doc, tr(lang, "languages"))
        _bullets(doc, language_strings(profile.languages))

    if profile.projects:
        _heading(doc, tr(lang, "projects"))
        for entry in profile.projects:
            _project_like(
                doc,
                getattr(entry, "title", str(entry)),
                getattr(entry, "organization", ""),
                getattr(entry, "period", ""),
                getattr(entry, "summary", ""),
            )

    if profile.certificates:
        _heading(doc, tr(lang, "certificates"))
        for entry in profile.certificates:
            _project_like(
                doc,
                getattr(entry, "title", str(entry)),
                getattr(entry, "issuer", ""),
                getattr(entry, "period", ""),
                getattr(entry, "summary", ""),
            )

    return doc

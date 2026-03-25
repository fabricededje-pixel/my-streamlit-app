from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from core.cv_formatters import language_strings, premium_contact_lines, skill_strings
from core.i18n import tr


def _set_layout(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _remove_table_borders(table):
    table_props = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    table_props.append(borders)


def _add_run(paragraph, text, *, bold=False, italic=False, size=10, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        from docx.shared import RGBColor

        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _header(doc, profile):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _remove_table_borders(table)

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.width = Cm(13.5)
    right.width = Cm(3.5)

    parts = profile.name.strip().split()
    first = parts[0] if parts else "Beispiel"
    rest = " ".join(parts[1:]) if len(parts) > 1 else "Name"

    paragraph_name = left.paragraphs[0]
    paragraph_name.paragraph_format.space_after = Pt(5)
    _add_run(paragraph_name, first + " ", size=26, color="374151")
    _add_run(paragraph_name, rest, bold=True, size=26, color="000000")

    if profile.job_title.strip():
        paragraph_job = left.add_paragraph()
        paragraph_job.paragraph_format.space_after = Pt(5)
        _add_run(paragraph_job, profile.job_title.strip(), size=11.5, color="4B5563")

    line1, line2 = premium_contact_lines(profile)
    if line1:
        paragraph_1 = left.add_paragraph()
        paragraph_1.paragraph_format.space_after = Pt(1)
        _add_run(paragraph_1, line1, size=10.5, color="374151")

    if line2:
        paragraph_2 = left.add_paragraph()
        paragraph_2.paragraph_format.space_after = Pt(0)
        _add_run(paragraph_2, line2, size=10.5, color="374151")

    if profile.photo_path:
        try:
            paragraph_image = right.paragraphs[0]
            paragraph_image.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_image = paragraph_image.add_run()
            run_image.add_picture(profile.photo_path, width=Cm(3.0))
        except Exception:
            pass


def _section_title(doc, title):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(4)

    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(15)

    paragraph_props = paragraph._element.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "6B7280")
    border.append(bottom)
    paragraph_props.append(border)


def _entry_block(doc, title, subtitle, details, period="", location=""):
    paragraph_top = doc.add_paragraph()
    paragraph_top.paragraph_format.space_after = Pt(2)

    run_title = paragraph_top.add_run(title.strip())
    run_title.bold = True
    run_title.font.name = "Arial"
    run_title.font.size = Pt(11.5)

    right_text = " | ".join([x for x in [period.strip(), location.strip()] if x])
    if right_text:
        run_meta = paragraph_top.add_run(f"    {right_text}")
        run_meta.italic = True
        run_meta.font.name = "Arial"
        run_meta.font.size = Pt(10)

    if subtitle:
        paragraph_sub = doc.add_paragraph()
        paragraph_sub.paragraph_format.space_after = Pt(4)
        run_sub = paragraph_sub.add_run(subtitle.upper())
        run_sub.font.name = "Arial"
        run_sub.font.size = Pt(9)

    for item in details:
        if str(item).strip():
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(str(item))
            run.font.name = "Arial"
            run.font.size = Pt(10)


def _summary_block(doc, text, lang):
    if not text.strip():
        return
    _section_title(doc, tr(lang, "summary"))
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def _skills_line(doc, label, values):
    if not values:
        return
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)

    run_label = paragraph.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.name = "Arial"
    run_label.font.size = Pt(10.5)

    run_values = paragraph.add_run(" | ".join(values))
    run_values.font.name = "Arial"
    run_values.font.size = Pt(10)


def build_german_premium_cv(profile):
    doc = Document()
    _set_layout(doc)

    lang = profile.language

    _header(doc, profile)

    if profile.summary.strip():
        _summary_block(doc, profile.summary.strip(), lang)

    if profile.experience:
        _section_title(doc, tr(lang, "experience"))
        for entry in profile.experience:
            _entry_block(
                doc,
                title=entry.title,
                subtitle=entry.company or "",
                details=entry.details,
                period=entry.period,
                location=entry.location,
            )

    if profile.education:
        _section_title(doc, tr(lang, "education"))
        for entry in profile.education:
            _entry_block(
                doc,
                title=entry.degree,
                subtitle=entry.school or "",
                details=entry.details,
                period=entry.period,
                location=entry.location,
            )

    if profile.projects:
        _section_title(doc, tr(lang, "projects"))
        for project in profile.projects:
            summary = str(getattr(project, "summary", "") or "").strip()
            _entry_block(
                doc,
                str(getattr(project, "title", project)),
                getattr(project, "organization", "") or "",
                [summary] if summary else [],
                getattr(project, "period", "") or "",
                "",
            )

    if profile.certificates:
        _section_title(doc, tr(lang, "certificates"))
        for certificate in profile.certificates:
            summary = str(getattr(certificate, "summary", "") or "").strip()
            _entry_block(
                doc,
                str(getattr(certificate, "title", certificate)),
                getattr(certificate, "issuer", "") or "",
                [summary] if summary else [],
                getattr(certificate, "period", "") or "",
                "",
            )

    if profile.languages or profile.skills:
        _section_title(doc, tr(lang, "skills"))

        langs = language_strings(profile.languages)
        if langs:
            _skills_line(doc, tr(lang, "languages"), langs)

        skills = skill_strings(profile.skills)
        if skills:
            _skills_line(doc, "Techstack", skills)

    return doc
